# luminai/ingest.py
import json, os, re, pathlib
from typing import List, Dict, Tuple
import numpy as np
import faiss
import fitz  # PyMuPDF
import pdfplumber
from pdf2image import convert_from_path
import pytesseract
from bs4 import BeautifulSoup
from .models import get_embedder
from docx import Document
import csv
import argparse
from pathlib import Path

HEADING_PAT = re.compile(r"(?m)^(#{1,3}\s+.+|[0-9]+\.[0-9.]*\s+.+)\s*$")
SUPPORTED_EXT = {".pdf", ".md", ".docx", ".txt", ".html", ".htm", ".csv"}
BASE_DIR = Path(__file__).resolve().parent.parent
INDEX_DIR = str(BASE_DIR / "indexes")
META_PATH = os.path.join(INDEX_DIR, "meta.json")
VEC_PATH  = os.path.join(INDEX_DIR, "vectors.npy")
FAISS_PATH= os.path.join(INDEX_DIR, "faiss.index")
CHUNK_SIZE = 600
CHUNK_OVERLAP = 120

def read_text(path: str) -> str:
    p = pathlib.Path(path)
    ext = p.suffix.lower()

    # ---------- PDF ----------
    if ext == ".pdf":
        text = ""

        # Try PyMuPDF (fast, good layout)
        try:
            with fitz.open(str(p)) as doc:
                text = "\n".join((page.get_text() or "") for page in doc)
        except Exception:
            text = text or ""

        # Fallback: pdfplumber (sometimes better on weird PDFs)
        if len(text.strip()) < 50:
            try:
                with pdfplumber.open(str(p)) as pdf:
                    text = "\n".join((page.extract_text() or "") for page in pdf.pages)
            except Exception:
                pass

        # Fallback: OCR (scanned/image-only PDFs)
        # Requires: pip install pdf2image pytesseract pillow
        # Note: On Windows, OCR also needs Tesseract installed and on PATH.
        if len(text.strip()) < 50:
            try:
                
                imgs = convert_from_path(str(p), dpi=200)
                text = "\n".join(pytesseract.image_to_string(im) for im in imgs)
            except Exception:
                pass

        return text or ""

    # ---------- DOCX ----------
    if ext == ".docx":
        try:
            d = Document(str(p))
            lines = []

            # Paragraphs
            for para in d.paragraphs:
                t = (para.text or "").strip()
                if t:
                    lines.append(t)

            # Tables -> pipe-delimited rows
            for table in d.tables:
                for row in table.rows:
                    cells = [(c.text or "").strip() for c in row.cells]
                    if any(cells):
                        lines.append(" | ".join(cells))

            return "\n".join(lines)
        except Exception:
            return ""

    # ---------- Markdown / Plain text ----------
    if ext in {".md", ".txt"}:
        try:
            return p.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return ""

    # ---------- (Optional) HTML ----------
    if ext in {".html", ".htm"}:
        # Requires: pip install beautifulsoup4 lxml
        try:
            html = p.read_text(encoding="utf-8", errors="ignore")
            soup = BeautifulSoup(html, "lxml")
            return soup.get_text(separator=" ")
        except Exception:
            return ""

    # ---------- (Optional) CSV ----------
    if ext == ".csv":
        try:
            out = []
            with p.open(newline="", encoding="utf-8", errors="ignore") as f:
                for row in csv.reader(f):
                    out.append(" | ".join((cell or "").strip() for cell in row))
            return "\n".join(out)
        except Exception:
            return ""

    # Unknown extension → ignore
    return ""

def clean_text(t: str) -> str:
    t = re.sub(r"\s+", " ", t).strip()
    return t

def split_by_headings(t: str) -> list[str]:
    parts, last = [], 0
    for m in HEADING_PAT.finditer(t):
        if m.start() > last:
            parts.append(t[last:m.start()])
        last = m.start()
    parts.append(t[last:])
    return [p.strip() for p in parts if p.strip()]

def chunk_text(t: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP):
    words = t.split()
    chunks, i = [], 0
    while i < len(words):
        chunk = words[i:i+chunk_size]
        if not chunk: break
        chunks.append(" ".join(chunk))
        i += max(1, chunk_size - overlap)
    return chunks

def find_files(root: str) -> List[str]:
    paths = []
    for r, _, files in os.walk(root):
        for f in files:
            if pathlib.Path(f).suffix.lower() in SUPPORTED_EXT:
                paths.append(os.path.join(r, f))
    return sorted(paths)

def build_index(data_dir: str,chunk_size: int = CHUNK_SIZE,overlap: int = CHUNK_OVERLAP,save_dir: str = INDEX_DIR,):
    # --- local helper: heading-aware splitter (Markdown and "1.2.3 ..." styles) ---
    HEADING_PAT = re.compile(r"(?m)^(#{1,3}\s+.+|[0-9]+(?:\.[0-9]+)*\s+.+)\s*$")

    def split_by_headings(text: str) -> List[Dict]:
        """
        Returns a list of sections: [{"heading": <str>, "body": <str>}]
        If no headings found, returns one section with empty heading.
        """
        sections = []
        last_pos = 0
        last_heading = ""
        matches = list(HEADING_PAT.finditer(text))
        if not matches:
            return [{"heading": "", "body": text.strip()}] if text.strip() else []

        for i, m in enumerate(matches):
            # body before the first heading (if any)
            if i == 0 and m.start() > 0:
                pre = text[:m.start()].strip()
                if pre:
                    sections.append({"heading": "", "body": pre})

            # current heading + its body up to next heading
            heading_line = m.group(0).strip()
            heading_text = re.sub(r"^#{1,3}\s+|\s+$", "", heading_line)  # strip leading #'s nicely
            next_start = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            body = text[m.end():next_start].strip()
            if body:
                sections.append({"heading": heading_text, "body": body})

        return sections

    # --- start ---
    os.makedirs(save_dir, exist_ok=True)

    files = find_files(data_dir)  # uses your existing helper
    texts: List[str] = []
    meta: List[Dict] = []

    for fp in files:
        raw = read_text(fp)              # robust extractor you just rewrote
        cleaned = clean_text(raw)        # your existing cleaner
        if not cleaned:
            continue

        sections = split_by_headings(cleaned) or [{"heading": "", "body": cleaned}]
        for s_idx, sec in enumerate(sections):
            body = sec["body"]
            heading = sec.get("heading", "")
            # chunk within each section for better topical cohesion
            for c_idx, ch in enumerate(chunk_text(body, chunk_size=chunk_size, overlap=overlap)):
                texts.append(ch)
                meta.append({
                    "file": fp,
                    "section_idx": s_idx,
                    "section_heading": heading,
                    "chunk_id": c_idx,
                })

    if not texts:
        raise ValueError(f"No ingestible text found under '{data_dir}'. "
                         f"Supported: {', '.join(sorted(SUPPORTED_EXT))}")

    # --- embeddings + FAISS (cosine via normalized vectors + inner product) ---
    embedder = get_embedder()
    vectors = embedder.encode(
        texts,
        show_progress_bar=True,
        convert_to_numpy=True,
        normalize_embeddings=True,  # important for IP-as-cosine
    )

    index = faiss.IndexFlatIP(vectors.shape[1])
    index.add(vectors)

    # --- persist artifacts ---
    vec_path   = os.path.join(save_dir, "vectors.npy")
    faiss_path = os.path.join(save_dir, "faiss.index")
    meta_path  = os.path.join(save_dir, "meta.json")

    os.makedirs(save_dir, exist_ok=True)
    vec_path, faiss_path, meta_path = _index_paths(save_dir)

    np.save(vec_path, vectors)
    faiss.write_index(index, faiss_path)
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump({"meta": meta, "texts": texts, "files": files}, f, ensure_ascii=False, indent=2)

    print(
        f"✅ Ingested {len(texts)} chunks from {len(files)} files "
        f"(dim={vectors.shape[1]}) → saved to '{save_dir}'"
    )
    return index, vectors, meta



def load_index(save_dir: str = INDEX_DIR):
    vec_path, faiss_path, meta_path = _index_paths(save_dir)
    vectors = np.load(vec_path)
    index = faiss.read_index(faiss_path)
    with open(meta_path, "r", encoding="utf-8") as f:
        meta = json.load(f)
    return index, vectors, meta

def _index_paths(save_dir: str = INDEX_DIR):
    vec_path   = os.path.join(save_dir, "vectors.npy")
    faiss_path = os.path.join(save_dir, "faiss.index")
    meta_path  = os.path.join(save_dir, "meta.json")
    return vec_path, faiss_path, meta_path

def index_ready(save_dir: str = INDEX_DIR) -> bool:
    vec_path, faiss_path, meta_path = _index_paths(save_dir)
    return all(os.path.exists(p) for p in (vec_path, faiss_path, meta_path))

if __name__ == "__main__":
    p = argparse.ArgumentParser()
    p.add_argument("data_dir")
    args = p.parse_args()
    build_index(args.data_dir)
    print("✅ Ingest complete. Index at", INDEX_DIR)
